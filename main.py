import os
import openai
from docx import Document

openai.api_key = "your_openai_api_key"

def translate_text(text, source_lang="pl", target_lang="en"):
    result = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Translate the following HTML-text from Polish to English as a native English speaker: {text}",
        max_tokens=50,
        n=1,
        stop=None,
        temperature=0.8,
    )

    return result.choices[0].text.strip()


def translate_docx(input_file, output_file):
    input_doc = Document(input_file)
    output_doc = Document()

    for paragraph in input_doc.paragraphs:
        translated_paragraph = output_doc.add_paragraph()
        for run in paragraph.runs:
            translated_text = translate_text(run.text)
            translated_run = translated_paragraph.add_run(translated_text)
            translated_run.bold = run.bold
            translated_run.italic = run.italic
            translated_run.underline = run.underline
            translated_run.style = run.style
            translated_run.font.name = run.font.name
            translated_run.font.size = run.font.size

    output_doc.save(output_file)


input_file = "input.docx"
output_file = "translated_output.docx"

translate_docx(input_file, output_file)
print("The translation is completed. The result saved in", output_file)
